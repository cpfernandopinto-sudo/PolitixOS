from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path("docs/Plano_Mestre_PolitixOS_2026.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
CYAN = RGBColor(0, 174, 239)
GRAY = RGBColor(90, 101, 115)
LIGHT = "E8EEF5"
DARK = "0B2545"
WHITE = RGBColor(255, 255, 255)
RED = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)
GREEN = RGBColor(20, 110, 75)


def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_ref)
        p._p.get_or_add_pPr().append(num_pr)
        p.add_run(item)


def add_callout(doc, label, text, fill="E8EEF5", color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " — ")
    font(r, 11, True, color)
    r = p.add_run(text)
    font(r, 11, False, color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        font(r, 9.5, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            font(r, 9.3, False, RGBColor(35, 45, 55))
    set_table_widths(table, widths)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(.492)
sec.footer_distance = Inches(.492)

# Standard Business Brief tokens
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

for list_name in ("List Bullet", "List Number"):
    st = doc.styles[list_name]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(.5)
    st.paragraph_format.first_line_indent = Inches(-.25)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.167

# Header/footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("POLITIXOS  |  PLANO MESTRE DE PRODUTO")
font(r, 8.5, True, GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run("Documento de trabalho • versão 1.0 • agosto de 2026")
font(r, 8.5, False, GRAY)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(52)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("POLITIXOS")
font(r, 11, True, CYAN)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Plano Mestre de Produto")
font(r, 28, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(22)
r = p.add_run("Estratégia, arquitetura, experiência, automação, inteligência e expansão comercial")
font(r, 14, False, GRAY)

add_callout(doc, "Objetivo", "Transformar o PolitixOS de um conjunto avançado de dashboards em uma plataforma confiável, comercializável e escalável de inteligência política multicanal.")

add_matrix(doc, ["Documento", "Escopo", "Horizonte", "Cadência"], [["Plano Mestre", "Produto completo", "22 semanas", "Sprints de 2 semanas"]], [1700, 3300, 1800, 2560])

doc.add_page_break()
add_heading(doc, "1. Sumário executivo", 1)
doc.add_paragraph("O PolitixOS já possui uma base visual consistente, módulos relevantes e capacidade real de captura e análise. A próxima evolução, porém, não deve ser conduzida por alterações isoladas de interface ou pela simples inclusão de novos canais. O produto precisa consolidar confiabilidade de dados, governança das automações, padronização analítica, rastreabilidade e posicionamento comercial.")
doc.add_paragraph("A ordem recomendada é corrigir a fundação operacional, estruturar a experiência executiva, normalizar relatórios e KPIs, validar a narrativa comercial e, então, ampliar canais e recursos avançados.")
add_callout(doc, "Posicionamento", "Uma central de inteligência política multicanal que transforma sinais públicos dispersos em prioridades, alertas e ações estratégicas fundamentadas em evidências.", "DDEBF7", NAVY)

add_heading(doc, "2. Diagnóstico atual", 1)
doc.add_paragraph("A análise das 29 telas e da arquitetura atual encontrou riscos objetivos que afetam confiança, operação e vendas.")
add_bullets(doc, [
    "Radar X exibindo NaN no termômetro e misturando dados antigos de candidatos diferentes.",
    "Fluxos do X aparecem sem webhooks configurados na tela de automação.",
    "Instagram possui posts coletados, mas não apresenta comentários analisados.",
    "Existem conteúdos marcados como sem análise, indicando pipeline incompleto.",
    "Todos os candidatos aparecem inativos, embora existam dados em vários painéis.",
    "Recomendações estratégicas aparecem repetidas ou genéricas.",
    "Investigação Profunda contém fontes repetidas e conteúdos inválidos com undefined/null.",
    "Alguns indicadores sugerem tempo real mesmo quando os dados estão desatualizados.",
    "Há KPIs semanticamente próximos sem fórmula, período e origem claramente informados.",
    "Tabelas e relatórios têm boa densidade, mas ainda precisam de hierarquia e leitura executiva padronizadas.",
])
add_callout(doc, "Risco comercial", "Uma demonstração visualmente forte pode perder credibilidade se o cliente encontrar números inválidos, dados antigos, etapas incompletas ou recomendações sem evidência.", "FDE9E7", RED)

add_heading(doc, "3. Princípios do produto", 1)
add_numbered(doc, [
    "Confiabilidade antes de expansão.",
    "Supabase como fonte oficial do estado do produto.",
    "n8n como orquestrador, não como banco de estado ou interface pública.",
    "Toda análise deve apontar suas evidências.",
    "Todo KPI deve ter fórmula, origem, período e data de atualização.",
    "IA recomenda e explica; decisões sensíveis permanecem revisáveis por pessoas.",
    "Nenhum canal novo entra em produção sem teste de viabilidade, custo e conformidade.",
    "A experiência deve priorizar ação e confiança, não apenas quantidade de gráficos.",
])

add_heading(doc, "4. Modelo de valor", 1)
add_matrix(doc, ["Pilar", "O que entrega"], [
    ["Capturar", "Coleta contínua de sinais em notícias, redes e canais autorizados."],
    ["Compreender", "Temas, sentimento, contexto, atores e evolução das narrativas."],
    ["Priorizar", "Risco, oportunidade, criticidade e urgência operacional."],
    ["Agir", "Protocolos, recomendações e próximos passos acionáveis."],
    ["Comprovar", "Fontes, histórico, atualização, confiança e auditoria."],
], [1800, 7560])

add_heading(doc, "5. Arquitetura recomendada: Supabase + n8n", 1)
doc.add_paragraph("A recomendação é manter o n8n no curto e médio prazo. Uma migração total agora acrescentaria custo e atraso sem resolver o principal problema. Entretanto, o papel do n8n precisa ser delimitado.")
add_heading(doc, "Responsabilidades do n8n", 2)
add_bullets(doc, ["Integração com APIs externas.", "Captura de publicações e comentários.", "Encadeamento e agendamento de etapas.", "Processamentos administrativos e chamadas a serviços de IA.", "Retentativas técnicas e transformação de payloads."])
add_heading(doc, "Responsabilidades do Supabase e da aplicação", 2)
add_bullets(doc, ["Fonte oficial do estado de trabalhos, execuções e etapas.", "Autenticação, autorização e isolamento por cliente.", "Histórico auditável, progresso, custos e métricas operacionais.", "Filas duráveis, checkpoints e deduplicação.", "Atualização da interface e contratos estáveis de dados."])
add_heading(doc, "Estruturas de controle", 2)
add_matrix(doc, ["Estrutura", "Finalidade"], [
    ["automation_jobs", "Demanda criada pelo usuário ou agenda."],
    ["automation_runs", "Cada execução do trabalho."],
    ["automation_steps", "Captura, comentários, IA e consolidação."],
    ["capture_checkpoints", "Cursor, data e último item processado."],
    ["source_health", "Saúde, latência e última atualização por canal."],
    ["ai_usage_ledger", "Modelo, tokens, custo, etapa e cliente."],
    ["content_versions", "Evolução de métricas e comentários."],
    ["automation_events", "Auditoria completa das transições."],
    ["alerts", "Alertas derivados, reconhecidos e resolvidos."],
], [2600, 6760])
add_callout(doc, "Decisão", "Manter o n8n, proteger todos os acionamentos no servidor e usar o Supabase como plano de controle. Migrar apenas trabalhadores de alto volume ou baixa latência quando as métricas justificarem.")

add_heading(doc, "6. Roadmap priorizado", 1)
add_matrix(doc, ["Prioridade", "Entregas"], [
    ["P0", "Qualidade dos dados, segurança, automação, KPIs, custos e observabilidade."],
    ["P1", "Design system, Visão Geral 2.0, Investigação Profunda 2.0, site e apresentação."],
    ["P2", "Cidade, análise por tema/hashtag, X, YouTube, Facebook e WhatsApp."],
    ["P3", "Chat com dados e verificação de alegações."],
    ["Experimento", "TikTok condicionado a acesso, atualização, conformidade e viabilidade."],
], [1600, 7760])

add_heading(doc, "7. Plano de sprints", 1)
sprints = [
    ("Sprint 0 — Auditoria e fundação", "Semanas 1–2", ["Inventário completo dos fluxos n8n e contratos de dados.", "Auditoria de canais, webhooks, credenciais e permissões.", "Glossário de KPIs com fórmula, origem, unidade e periodicidade.", "Linha de base de custos de IA e APIs.", "Auditoria UX por página e mapa de bugs/inconsistências.", "Definição de personas, casos de uso e SLAs por canal.", "Separação planejada entre desenvolvimento e produção do n8n."], "Sabemos o que funciona, o que falha, quanto custa e quando cada dado foi atualizado."),
    ("Sprint 1 — Centro de controle da automação", "Semanas 3–4", ["Registrar trabalhos, execuções e etapas no Supabase.", "Mover acionamentos para rotas autenticadas no servidor.", "Criar identificadores únicos, deduplicação e idempotência.", "Implementar progresso, cancelamento, retentativas e reprocessamento.", "Registrar erros e saúde das fontes.", "Reconstruir a tela de automação com estado real."], "Uma execução pode ser acompanhada, retomada e auditada do início ao fim."),
    ("Sprint 2 — Design system e KPIs", "Semanas 5–6", ["Hierarquia de títulos, cards, gráficos, tabelas e modais.", "Estados de carregamento, ausência de dados e falha.", "Componentes de atualização, cobertura e confiança.", "Padrão de cores semânticas e responsividade.", "Revisão de todos os KPIs das páginas atuais."], "Todas as páginas seguem o mesmo contrato visual e informacional."),
    ("Sprint 3 — Visão Geral 2.0", "Semanas 7–8", ["Briefing diário da Politix IA no início.", "KPIs executivos com cobertura e atualização.", "Mudanças, riscos e oportunidades prioritárias.", "Distribuição por canal e próximas ações.", "Evidências clicáveis para cada recomendação."], "A primeira tela resume o dia e leva o usuário diretamente ao que exige decisão."),
    ("Sprint 4 — Automação adaptativa", "Semanas 9–10", ["Encadear publicação, comentários, IA, consolidação e alertas.", "Criar checkpoints e recaptura de novos comentários.", "Variar frequência conforme idade e velocidade do conteúdo.", "Exibir progresso e volume real por etapa."], "Posts recentes são revisitados sem duplicação e comentários novos entram na análise."),
    ("Sprint 5 — Investigação Profunda 2.0", "Semanas 11–12", ["Padronizar resumo executivo, fatos, alegações, linha do tempo, atores e narrativas.", "Separar riscos, oportunidades, recomendações, fontes, confiança e limitações.", "Validar fontes duplicadas e impedir publicação com conteúdo inválido.", "Manter relatório técnico bruto em área secundária."], "Nenhum dossiê é publicado com fonte inválida, ausência de evidência ou campos técnicos quebrados."),
    ("Sprint 6 — Cidade e narrativas", "Semanas 13–14", ["Página de inteligência territorial.", "Temas, fontes locais, atores, riscos, demandas e oportunidades.", "Explorador por hashtag, palavra, pessoa e tema.", "Evolução, canais, influenciadores, candidatos e conteúdos impulsionadores."], "O produto permite compreender uma narrativa por território e entre canais."),
    ("Sprint 7 — X e YouTube", "Semanas 15–16", ["Prova de conceito da API do X com orçamento controlado.", "Cursores incrementais, deduplicação e custo por candidato/dia.", "Correção de datas, candidatos misturados, NaN e webhooks.", "YouTube: canais, vídeos, busca, comentários, engajamento e análise."], "Cada canal possui custo, cobertura, atualização e qualidade mensurados."),
    ("Sprint 8 — Facebook, WhatsApp e estudo TikTok", "Semanas 17–18", ["Facebook: validar permissões antes da página definitiva.", "WhatsApp: recepção consentida, classificação, anonimização e encaminhamento.", "Resumo diário e alertas de crescimento de temas.", "TikTok: estudo de acesso, atraso, dados disponíveis e uso comercial."], "Somente integrações permitidas, sustentáveis e observáveis avançam para produção."),
    ("Sprint 9 — Chat com dados", "Semanas 19–20", ["Perguntas em linguagem natural.", "Respostas limitadas aos dados autorizados.", "Citações clicáveis e filtros por candidato, canal, cidade e período.", "Registro de perguntas, atualização e cobertura.", "Resposta explícita quando faltarem evidências."], "Toda resposta é rastreável e respeita o acesso do usuário."),
    ("Sprint 10 — Verificação de alegações", "Semanas 21–22", ["Extrair alegações verificáveis e seus autores.", "Buscar fontes primárias e comparar evidências.", "Classificar resultado, confiança e limitações.", "Criar revisão humana para casos sensíveis."], "Não existe veredito binário sem evidência, contexto e possibilidade de revisão."),
]
for title, period, items, done in sprints:
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    r = p.add_run(period)
    font(r, 10, True, GRAY)
    add_bullets(doc, items)
    add_callout(doc, "Critério de conclusão", done, "F4F6F9", NAVY)

add_heading(doc, "8. Visão Geral 2.0", 1)
doc.add_paragraph("A página inicial deve funcionar como um briefing executivo diário, reduzindo a necessidade de percorrer todo o sistema para descobrir o que mudou.")
add_callout(doc, "Exemplo", "Bom dia. Nas últimas 24 horas, o PolitixOS analisou 1.284 sinais relacionados aos candidatos monitorados. Foram identificados 12 alertas prioritários, 3 mudanças relevantes de narrativa e 2 oportunidades de posicionamento.", "E5F7FB", NAVY)
add_numbered(doc, ["Briefing diário gerado por IA.", "Atualização, cobertura e confiabilidade.", "KPIs executivos.", "Principais mudanças desde o período anterior.", "Riscos prioritários.", "Oportunidades prioritárias.", "Distribuição por canal.", "Próximas ações recomendadas."])
doc.add_paragraph("A IA deverá citar os dados que sustentam suas conclusões. Indicadores semelhantes, como temperatura, risco, estado político e termômetro, só devem coexistir se houver definições e papéis distintos.")

add_heading(doc, "9. Automação inteligente e recaptura", 1)
add_numbered(doc, ["Captar publicações.", "Salvar e deduplicar.", "Captar comentários.", "Analisar conteúdo.", "Consolidar KPIs.", "Gerar alertas.", "Agendar nova varredura."])
add_matrix(doc, ["Momento inicial", "Objetivo"], [
    ["Imediato", "Registrar publicação e métricas iniciais."],
    ["+1 hora", "Captar primeiras respostas e medir aceleração."],
    ["+6 horas", "Reavaliar risco e narrativa."],
    ["+24 horas", "Consolidar primeiro ciclo."],
    ["+72 horas", "Captar repercussão tardia."],
    ["+7 dias", "Fechar ciclo ou manter se houver atividade."],
], [2100, 7260])
doc.add_paragraph("Os intervalos são hipóteses iniciais. A frequência deverá aumentar quando houver aceleração de comentários e diminuir quando não houver alterações.")

add_heading(doc, "10. Inteligência territorial e análise por tema", 1)
add_heading(doc, "Página Cidade", 2)
add_bullets(doc, ["Resumo político da cidade.", "Temas dominantes e percepção por tema.", "Principais fontes locais e atores políticos.", "Bairros ou regiões mencionadas.", "Riscos emergentes, demandas e dores locais.", "Comparação entre candidatos.", "Oportunidades de agenda e posicionamento."])
add_heading(doc, "Explorador de narrativa", 2)
add_bullets(doc, ["Busca por hashtag, palavra, pessoa ou tema.", "Variações e termos relacionados.", "Volume, evolução, sentimento e risco.", "Plataformas, autores e fontes influentes.", "Candidatos relacionados e conteúdos que impulsionaram a narrativa.", "Recomendação estratégica sustentada por evidências."])

add_heading(doc, "11. Expansão por canal", 1)
add_matrix(doc, ["Canal", "Direção", "Condição para avançar"], [
    ["X", "Prova de conceito com poucos candidatos.", "Custo, deduplicação, atualização e qualidade medidos."],
    ["YouTube", "Canais, vídeos, comentários e análise.", "Quota e captura incremental controladas."],
    ["Facebook", "Páginas oficiais e conteúdo permitido.", "Permissões e revisão do aplicativo validadas."],
    ["WhatsApp", "Inteligência das mensagens recebidas.", "Consentimento, retenção e proteção de dados definidos."],
    ["TikTok", "Experimento de viabilidade.", "Acesso, atraso, conformidade e uso comercial aprovados."],
], [1400, 3500, 4460])

add_heading(doc, "12. Economia de IA", 1)
doc.add_paragraph("Antes de trocar modelos ou fornecedores, o produto precisa registrar o custo real de cada chamada.")
add_bullets(doc, ["Cliente, candidato, canal e etapa.", "Modelo e versão do prompt.", "Tokens de entrada e saída.", "Custo estimado e tempo de resposta.", "Primeira análise ou reprocessamento.", "Resultado, erro e qualidade observada."])
add_heading(doc, "Medidas de economia", 2)
add_bullets(doc, ["Deduplicar antes de analisar.", "Não reanalisar conteúdo inalterado.", "Usar modelos menores para classificação simples.", "Reservar modelos maiores para riscos e relatórios.", "Processar em lotes e resumir grupos de comentários.", "Armazenar resultados pelo hash do conteúdo.", "Reprocessar somente quando conteúdo, prompt ou modelo mudar.", "Definir orçamento por cliente, canal e período."])

add_heading(doc, "13. Chat e verificação de alegações", 1)
add_heading(doc, "Chat com os dados", 2)
add_bullets(doc, ["Respostas baseadas apenas nos dados autorizados.", "Citações clicáveis.", "Filtros por candidato, cidade, canal e período.", "Controle de acesso e registro das perguntas.", "Informação de atualização e cobertura.", "Recusa explícita quando não houver evidência suficiente."])
add_heading(doc, "Verificação de alegações", 2)
doc.add_paragraph("O produto deve evitar o rótulo binário 'fake news'. A abordagem recomendada é extrair alegações verificáveis, localizar fontes primárias, comparar evidências, indicar confiança e encaminhar casos sensíveis para revisão humana.")
add_bullets(doc, ["Confirmada.", "Parcialmente confirmada.", "Sem evidências suficientes.", "Enganosa ou fora de contexto.", "Contradita por evidências.", "Em apuração."])

add_heading(doc, "14. Métricas do próprio produto", 1)
add_matrix(doc, ["Dimensão", "Indicadores"], [
    ["Operação", "Sucesso das automações, tempo de captura, falhas e retentativas."],
    ["Dados", "Atualidade, cobertura, duplicidade e percentual analisado."],
    ["Comentários", "Percentual capturado e atraso entre publicação e análise."],
    ["Custos", "API por candidato/dia e IA por mil itens."],
    ["Qualidade", "Alertas úteis, relatórios válidos e recomendações com evidência."],
    ["Produto", "Uso semanal, tempo até insight e ações realizadas."],
], [1800, 7560])

add_heading(doc, "15. Apresentação comercial", 1)
add_numbered(doc, ["A complexidade da inteligência política digital.", "O problema do monitoramento fragmentado.", "O que é o PolitixOS.", "Como sinais se transformam em ação.", "Visão Geral e briefing diário de IA.", "Radar de Notícias e redes sociais.", "Alertas, riscos e oportunidades.", "Investigação Profunda.", "Inteligência territorial e de narrativas.", "Governança, segurança e rastreabilidade.", "Roadmap multicanal.", "Demonstração e proposta comercial."])
add_callout(doc, "Regra comercial", "Não apresentar números fictícios de economia, precisão ou velocidade. Até termos medições, comunicar capacidades, fluxo de trabalho e ganhos operacionais esperados.", "FFF4CE", GOLD)

add_heading(doc, "16. Site comercial", 1)
add_numbered(doc, ["Hero: Inteligência política que transforma sinais em decisões.", "Problema do mercado.", "Como funciona.", "Módulos da plataforma.", "Casos de uso.", "Briefing de IA.", "Monitoramento multicanal.", "Investigação e evidências.", "Segurança e governança.", "Solicitação de demonstração."])
add_heading(doc, "Públicos iniciais", 2)
add_bullets(doc, ["Campanhas eleitorais.", "Mandatos e partidos.", "Agências de comunicação política.", "Relações governamentais e assuntos públicos.", "Salas de situação e gestão de crise."])

add_heading(doc, "17. Cronograma consolidado", 1)
add_matrix(doc, ["Semanas", "Entrega principal"], [
    ["1–2", "Auditoria e fundação."], ["3–4", "Automação e segurança."],
    ["5–6", "Design system e KPIs."], ["7–8", "Visão Geral 2.0."],
    ["9–10", "Automação adaptativa."], ["11–12", "Investigação Profunda 2.0."],
    ["13–14", "Cidade e narrativas."], ["15–16", "X e YouTube."],
    ["17–18", "Facebook, WhatsApp e estudo TikTok."], ["19–20", "Chat com dados."],
    ["21–22", "Verificação de alegações."],
], [1700, 7660])
doc.add_paragraph("A apresentação e o site podem começar no Sprint 1 e ser publicados após o Sprint 3, quando o produto já tiver narrativa comercial coerente, métricas definidas e telas confiáveis.")

add_heading(doc, "18. Critérios gerais de pronto", 1)
add_bullets(doc, ["Dados têm origem, período e atualização visíveis.", "Erros e estados vazios são tratados.", "Nenhum fluxo depende de webhook público no navegador.", "Execuções são rastreáveis e idempotentes.", "KPIs têm fórmula documentada.", "Recomendações abrem suas evidências.", "Custos de API e IA são registrados.", "A interface funciona em notebook e desktop.", "Funcionalidades sensíveis possuem revisão humana e auditoria.", "Documentação operacional é atualizada junto com a entrega."])

add_heading(doc, "19. Próximo passo", 1)
add_callout(doc, "Início recomendado", "Executar o Sprint 0 antes de qualquer nova página ou integração. O resultado será o diagnóstico técnico e operacional, o dicionário de KPIs, o mapa de automações, a linha de base de custos e o backlog reestimado para execução.", "E5F7FB", NAVY)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
