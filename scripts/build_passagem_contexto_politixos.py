from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path("docs/Passagem_de_Contexto_PolitixOS.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
CYAN = RGBColor(0, 174, 239)
GRAY = RGBColor(90, 101, 115)
INK = RGBColor(35, 45, 55)
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label} — ")
    set_font(run, 11, True, NAVY)
    run = paragraph.add_run(text)
    set_font(run, 11, False, NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_font(run, 9.5, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_font(run, 9.3, False, INK)
    set_table_widths(table, widths)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Compact Reference Guide preset
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.text = "POLITIXOS  |  PASSAGEM DE CONTEXTO"
set_font(header.runs[0], 9, True, GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.text = "Documento de continuidade • agosto de 2026"
set_font(footer.runs[0], 9, False, GRAY)

# Opening block: memo masthead
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(5)
run = kicker.add_run("DOCUMENTO DE CONTINUIDADE DO PROJETO")
set_font(run, 10, True, CYAN)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Passagem de Contexto — PolitixOS")
set_font(run, 25, True, NAVY)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("Resumo operacional para iniciar uma nova conversa sem perder decisões, prioridades e método de trabalho.")
set_font(run, 12.5, False, GRAY)

add_matrix(doc, ["Campo", "Informação"], [
    ("Produto", "PolitixOS — plataforma de inteligência política multicanal."),
    ("Responsável", "Fernando Pinto."),
    ("Finalidade", "Continuidade estratégica, acompanhamento de sprints e preparação de prompts para agentes de IA."),
    ("Documento-base", "Plano Mestre do PolitixOS 2026, versão 1.0."),
    ("Estado do handoff", "Planejamento consolidado; execução futura deve começar por auditoria e confirmação do estado real."),
], [1900, 7460])

add_callout(doc, "Instrução principal", "Use este documento junto com o Plano Mestre. Não altere código, dados, automações ou produção sem autorização explícita de Fernando.", "DDF3FA")

add_heading(doc, "1. Como este novo chat deve funcionar", 1)
add_bullets(doc, [
    "Ser a central estratégica do projeto: discutir, analisar resultados, registrar decisões e acompanhar o passo a passo.",
    "Criar prompts específicos para Claude, Codex e Antigravity, escolhendo o agente conforme a natureza da tarefa.",
    "Separar planejamento de execução. Ideias e análises não significam autorização para editar código ou publicar.",
    "Usar evidências: capturas de tela, dados reais, estrutura do banco, fluxos n8n, logs, custos e testes.",
    "Atualizar o Plano Mestre e este handoff quando decisões relevantes forem aprovadas.",
])

add_heading(doc, "2. Visão do produto", 1)
doc.add_paragraph(
    "O PolitixOS deve evoluir de um conjunto de painéis para um sistema operacional de inteligência política. "
    "O produto precisa transformar sinais de notícias e redes sociais em leitura executiva, alertas rastreáveis, "
    "investigações compreensíveis e ações recomendadas por IA, sempre sustentadas por evidências."
)
add_bullets(doc, [
    "Públicos iniciais: campanhas, mandatos, partidos, agências de comunicação, relações governamentais e salas de crise.",
    "Diferencial: consolidar monitoramento multicanal, risco, narrativa, território e recomendação em uma única operação.",
    "Princípio comercial: não prometer precisão, economia ou velocidade sem medição real.",
    "Princípio de IA: recomendações devem mostrar os dados que sustentam a conclusão e admitir insuficiência de evidência.",
])

add_heading(doc, "3. O que já foi discutido e ajustado em UX", 1)
add_heading(doc, "Visão Geral", 2)
add_bullets(doc, [
    "Foi construída uma linguagem visual corporativa escura, com cards azul-marinho, hierarquia mais compacta e uso controlado das cores semânticas.",
    "KPIs do topo foram compactados para ampliar a visão útil acima da dobra.",
    "Termômetro de crise e Alertas Prioritários foram redistribuídos para reduzir áreas vazias e melhorar leitura.",
    "A seção Leitura Analítica passou a organizar quatro gráficos em linha no notebook, preservando o radar multidimensional de Distribuição por Canal.",
    "Títulos e seções passaram por tentativas de padronização; uma versão excessivamente compacta foi rejeitada e revertida.",
])

add_heading(doc, "Radar de Notícias", 2)
add_bullets(doc, [
    "Cards principais foram reorganizados entre Termômetro de Crise, Status em Tempo Real e Feed Crítico.",
    "Contornos brancos foram rejeitados por não pertencerem ao sistema visual.",
    "Leitura Analítica aprovada conceitualmente em duas linhas: três cards na primeira e dois gráficos na segunda.",
    "Primeira linha: Principais Temas Negativos, Fontes com Maior Impacto e Distribuição do Impacto.",
    "Segunda linha: Evolução do Risco e Linha do Tempo de Crise — Últimas 24 horas.",
    "Filtros Top 5/Top 10 precisam funcionar; a preferência atual é Top 10 como padrão, com barras mais estreitas e maior área para rótulos.",
    "A Base Completa de Monitoramento ganhou a coluna Candidato; o espaçamento entre Fonte e Candidato deve permanecer compacto.",
])

add_heading(doc, "Detalhes e modais", 2)
add_bullets(doc, [
    "Notícias devem abrir um modal com os dados reais do registro selecionado, não conteúdo genérico.",
    "O modal de notícia precisa incluir a ação sugerida pela IA, considerada elemento central da proposta de valor.",
    "Modais de Instagram e notícias não devem criar grandes áreas vazias; usar altura máxima de viewport, rolagem interna e posicionamento estável.",
    "Alertas e contagens precisam representar os registros reais, sem duplicação de informações.",
])

add_heading(doc, "4. Regra de ouro para confirmar o estado atual", 1)
add_callout(doc, "Atenção", "As conversas incluíram ajustes locais, solicitações de deploy e novas iterações posteriores. Antes de planejar qualquer mudança, comparar repositório, branch, produção e automações. Não presumir que todos estão sincronizados.", "FFF2CC")
add_matrix(doc, ["Verificação", "Pergunta que deve ser respondida"], [
    ("Código", "Qual branch e commit representam a versão mais recente aprovada?"),
    ("Produção", "Qual deploy está ativo em app.politixos.ia.br e ele corresponde ao repositório?"),
    ("Banco", "Quais tabelas, campos e políticas alimentam cada página e modal?"),
    ("n8n", "Quais fluxos estão ativos, com que agenda, credenciais, retentativas e dependências?"),
    ("APIs", "Quais fontes estão contratadas, quais limites existem e quanto cada canal custa?"),
    ("IA", "Quais modelos, prompts, tokens, custos e versões estão em uso?"),
], [1900, 7460])

add_heading(doc, "5. Direção técnica e papel do n8n", 1)
doc.add_paragraph(
    "O n8n deve permanecer inicialmente como orquestrador, porque já concentra parte das automações. "
    "A decisão entre manter, migrar ou adotar arquitetura híbrida só deve ocorrer depois de uma auditoria baseada em confiabilidade, custo e observabilidade."
)
add_bullets(doc, [
    "Supabase: fonte de verdade para configurações, execuções, checkpoints, conteúdo bruto, análises e custos.",
    "n8n: coordenação de etapas, agendas, webhooks e integrações enquanto for confiável e sustentável.",
    "Aplicação: configura buscas, acompanha progresso, mostra resultados e permite reprocessamento controlado.",
    "Requisitos mínimos: idempotência, deduplicação, checkpoints, retentativas, logs, alertas e rastreabilidade.",
    "Fluxo esperado: captar publicação → salvar/deduplicar → captar comentários → analisar → consolidar KPIs → gerar alertas → agendar nova varredura.",
    "Recaptura deve ser adaptativa: maior frequência para conteúdo novo ou acelerando; redução progressiva sem alterações; encerramento controlado.",
])

add_heading(doc, "6. Backlog estratégico consolidado", 1)
add_matrix(doc, ["Frente", "Escopo resumido", "Prioridade"], [
    ("Fundação", "Auditoria de dados, produção, n8n, APIs, IA, segurança e observabilidade.", "Imediata"),
    ("Visão Geral 2.0", "Briefing diário de IA, mudanças relevantes, riscos, oportunidades e cobertura.", "Alta"),
    ("Automação", "Tela de configuração, progresso, encadeamento e recaptura adaptativa.", "Alta"),
    ("Relatórios", "Padronizar Investigação Profunda e destacar ações sugeridas com evidências.", "Alta"),
    ("Narrativas", "Busca por hashtag, termo, pessoa e tema em múltiplas páginas e canais.", "Alta"),
    ("Território", "Página Cidade com contexto político, temas, fontes, bairros e oportunidades.", "Média"),
    ("Canais", "Avaliar X; criar YouTube, Facebook e WhatsApp; TikTok como experimento posterior.", "Faseada"),
    ("IA e custos", "Telemetria, cache, deduplicação, roteamento de modelos e orçamento.", "Transversal"),
    ("Chat com dados", "Perguntas com citações, filtros, controle de acesso e cobertura explícita.", "Após fundação"),
    ("Verificação", "Análise de alegações e evidências, evitando rótulo binário de fake news.", "Após fundação"),
    ("Comercial", "Apresentação e site para começar vendas com demonstração confiável.", "Paralela"),
], [1700, 5960, 1700])

add_heading(doc, "7. Roadmap de referência", 1)
add_matrix(doc, ["Sprint", "Objetivo", "Resultado esperado"], [
    ("0", "Auditoria e baseline", "Mapa real da arquitetura, integrações, dados, custos e riscos."),
    ("1", "Fundação de dados", "Contratos, deduplicação, rastreabilidade e indicadores de cobertura."),
    ("2", "Automação observável", "Execuções com status, checkpoints, retentativas e logs."),
    ("3", "Automação adaptativa", "Recaptura de comentários e encadeamento controlado."),
    ("4", "UX e design system", "Padrões de cards, títulos, tabelas, modais e responsividade."),
    ("5", "Visão Geral 2.0", "Briefing executivo de IA e KPIs confiáveis."),
    ("6", "Investigação e narrativas", "Relatórios padronizados e explorador por tema/hashtag."),
    ("7", "Cidade", "Inteligência territorial e comparação política local."),
    ("8", "Expansão de canais", "Provas controladas de X, YouTube, Facebook e WhatsApp."),
    ("9", "Chat com dados", "Respostas rastreáveis, autorizadas e citadas."),
    ("10", "Verificação de alegações", "Evidências, confiança, limitações e revisão humana."),
], [1000, 2920, 5440])

add_heading(doc, "8. Método obrigatório para cada demanda", 1)
add_bullets(doc, [
    "Problema e objetivo: explicar o que está errado e qual resultado de negócio é esperado.",
    "Diagnóstico: inspecionar dados, tela, código e automação sem alterar nada.",
    "Proposta: descrever solução, alternativas, riscos e impacto.",
    "Agente: indicar Claude, Codex ou Antigravity e justificar a escolha.",
    "Prompt: delimitar arquivos, comportamento, dados reais, restrições e o que não deve ser tocado.",
    "Critérios de aceite: definir verificações funcionais, visuais, de dados e responsividade.",
    "Execução e testes: implementar somente após autorização e apresentar evidências.",
    "Registro: documentar resultado, pendências, decisão e eventual deploy.",
])

add_heading(doc, "9. Tipos de prompt a produzir", 1)
add_matrix(doc, ["Tipo", "Permissão"], [
    ("Diagnóstico", "Somente leitura e análise; não altera arquivos, banco ou produção."),
    ("Implementação", "Autoriza mudanças específicas, delimitadas por escopo e arquivos."),
    ("Correção", "Resolve defeitos encontrados na validação sem redesenhar o restante."),
    ("Auditoria", "Verifica segurança, dados, custos, desempenho e qualidade."),
    ("Deploy", "Publica somente a versão aprovada, com testes e plano de retorno."),
], [1900, 7460])

add_heading(doc, "10. Primeiro passo recomendado", 1)
add_callout(doc, "Início do processo", "Executar o Sprint 0 em modo diagnóstico. O resultado deve ser um inventário verificável e um backlog priorizado — ainda sem reformulação ampla do código.", "DDF3FA")
add_bullets(doc, [
    "Mapear repositório, tecnologias, ambientes, branches e deploy ativo.",
    "Inventariar tabelas, relacionamentos, políticas e origem dos dados exibidos.",
    "Inventariar fluxos n8n, agendas, falhas, credenciais, volumes e dependências.",
    "Mapear APIs por canal, limites, custos, termos e lacunas.",
    "Medir uso de IA por tarefa, prompt, modelo, tokens, custo e reprocessamento.",
    "Conferir página por página: objetivo, KPIs, dados duplicados, vazios, cortes e inconsistências.",
    "Entregar riscos, quick wins, dependências e proposta final de sprints.",
])

add_heading(doc, "11. Mensagem pronta para abrir o novo chat", 1)
message = (
    "Este chat será a central estratégica do projeto PolitixOS. Leia integralmente os dois documentos anexados: "
    "Plano Mestre do PolitixOS 2026 e Passagem de Contexto do PolitixOS. Use-os como contexto do projeto, mas confirme "
    "o estado atual do repositório, da produção, do banco e das automações antes de assumir que uma alteração está ativa. "
    "Aqui vamos discutir produto, UX, conteúdo, dados, n8n, APIs, IA, custos, roadmap, sprints e resultados. Não execute "
    "alterações no código, banco, automações ou produção sem minha autorização explícita. Para cada demanda, organize: "
    "objetivo, diagnóstico, proposta, riscos, agente indicado, prompt de execução, critérios de aceite, testes e registro "
    "da decisão. O primeiro trabalho deverá ser planejar o Sprint 0 — Auditoria e Fundação."
)
add_callout(doc, "Copiar e enviar", message, "E8EEF5")

add_heading(doc, "12. Arquivos de referência", 1)
add_bullets(doc, [
    "Plano Mestre do PolitixOS 2026: visão completa, arquitetura, backlog, sprints e direção comercial.",
    "Passagem de Contexto do PolitixOS: decisões operacionais, estado conhecido, regras e ponto de retomada.",
    "politixos.zip: material de referência fornecido; deve ser comparado com o repositório antes de ser tratado como versão vigente.",
    "Capturas de tela: evidências históricas da evolução de UX; novas decisões devem usar capturas atualizadas.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
